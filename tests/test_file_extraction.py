"""Unit tests for ui/file_extraction.py.

Fixtures are built in-memory rather than committed as binary blobs, so
what each test actually contains stays readable as Python, not an opaque
checked-in file. The PDF fixture is hand-built at the byte level (no
reportlab/fpdf dependency just for a test) - a minimal but genuinely
valid PDF with real extractable text.

The OCR tests run real Tesseract, not a mock - this machine has it
installed (see the root README's UI-redesign section for the local
install note), so this is a real, not hoped-for, verification of the
image -> OCR -> text pipeline. The "Tesseract not available" path is
separately verified via mocking, since forcing a real absence isn't
practical in an environment where it's actually installed.
"""

from __future__ import annotations

import io

import pytest

from ui.file_extraction import (
    ImageDecodeError,
    TesseractNotAvailableError,
    UnsupportedFileTypeError,
    extract_text_from_docx,
    extract_text_from_image,
    extract_text_from_pdf,
    extract_text_from_txt,
    extract_text_from_upload,
    is_tesseract_available,
)


def _build_minimal_pdf(text: str) -> bytes:
    """Hand-build a minimal, valid, single-page PDF containing `text` as
    real extractable content - see this file's module docstring."""
    stream_content = f"BT /F1 14 Tf 10 100 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 300 144] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream_content)).encode() + b" >>\nstream\n" + stream_content + b"\nendstream",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        b"trailer\n"
        + f"<< /Size {len(objects) + 1} /Root 1 0 R >>\n".encode()
        + b"startxref\n"
        + f"{xref_offset}\n".encode()
        + b"%%EOF"
    )
    return bytes(out)


def _build_docx(paragraphs: list[str], table_rows: list[tuple[str, str]] | None = None) -> bytes:
    import docx

    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        table = d.add_table(rows=0, cols=2)
        for a, b in table_rows:
            row = table.add_row()
            row.cells[0].text = a
            row.cells[1].text = b
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _build_text_image(text: str) -> bytes:
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (600, 120), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 40), text, fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


# --- Plain text ---


def test_extract_text_from_txt_decodes_utf8():
    assert extract_text_from_txt("A sterile hip implant.".encode("utf-8")) == "A sterile hip implant."


def test_extract_text_from_txt_replaces_undecodable_bytes_instead_of_crashing():
    result = extract_text_from_txt(b"\xff\xfe not valid utf-8 on its own")
    assert isinstance(result, str)  # must not raise


# --- PDF ---


def test_extract_text_from_pdf_real_fixture():
    pdf_bytes = _build_minimal_pdf("A sterile hip implant with a titanium coating.")
    assert extract_text_from_pdf(pdf_bytes) == "A sterile hip implant with a titanium coating."


# --- DOCX ---


def test_extract_text_from_docx_paragraphs():
    docx_bytes = _build_docx(["A cerebrovascular stent implanted in the MCA."])
    assert extract_text_from_docx(docx_bytes) == "A cerebrovascular stent implanted in the MCA."


def test_extract_text_from_docx_includes_table_cells():
    """Technical files often put key specs (material, sterility) in
    tables rather than prose - these must not be silently dropped."""
    docx_bytes = _build_docx(["Device summary"], table_rows=[("Sterility", "Sterile, single use")])
    result = extract_text_from_docx(docx_bytes)
    assert "Device summary" in result
    assert "Sterility" in result
    assert "Sterile, single use" in result


# --- Image / OCR (real Tesseract, not mocked) ---


def test_tesseract_is_available_in_this_environment():
    """Sanity check for the test environment itself - if this fails, the
    OCR tests below aren't testing what they claim to."""
    assert is_tesseract_available() is True


def test_is_tesseract_available_returns_false_when_not_found(monkeypatch):
    import pytesseract

    def _raise(*args, **kwargs):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(pytesseract, "get_tesseract_version", _raise)
    assert is_tesseract_available() is False


def test_extract_text_from_image_real_ocr():
    image_bytes = _build_text_image("Implantable cardiac pacemaker")
    result = extract_text_from_image(image_bytes)
    assert "pacemaker" in result.lower()


def test_extract_text_from_image_no_text_yields_empty_string():
    """A bare photo of a device with no visible text should honestly
    yield nothing - not a hallucinated guess. This is the documented
    limit of OCR-only image reading (see the module docstring)."""
    from PIL import Image

    blank = Image.new("RGB", (200, 200), color="white")
    buf = io.BytesIO()
    blank.save(buf, format="PNG")
    assert extract_text_from_image(buf.getvalue()) == ""


def test_extract_text_from_image_rejects_corrupt_image():
    with pytest.raises(ImageDecodeError):
        extract_text_from_image(b"this is not an image file")


def test_extract_text_from_image_wraps_tesseract_not_found(monkeypatch):
    import pytesseract

    def _raise(*args, **kwargs):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(pytesseract, "image_to_string", _raise)
    image_bytes = _build_text_image("Implantable cardiac pacemaker")
    with pytest.raises(TesseractNotAvailableError):
        extract_text_from_image(image_bytes)


# --- Dispatcher ---


def test_extract_text_from_upload_dispatches_by_extension():
    assert extract_text_from_upload("device.txt", b"hello") == "hello"
    assert "hip implant" in extract_text_from_upload(
        "spec.pdf", _build_minimal_pdf("A hip implant.")
    ).lower()


def test_extract_text_from_upload_rejects_unsupported_extension():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload("device.xlsx", b"whatever")


def test_extract_text_from_upload_rejects_extensionless_filename():
    with pytest.raises(UnsupportedFileTypeError):
        extract_text_from_upload("noextension", b"whatever")


# --- Truncation ---


def test_extracted_text_is_truncated_past_the_length_cap():
    from ui.file_extraction import MAX_EXTRACTED_CHARS

    huge = "A" * (MAX_EXTRACTED_CHARS + 500)
    result = extract_text_from_txt(huge.encode())
    assert len(result) <= MAX_EXTRACTED_CHARS + len("\n[... truncated - text was longer than the limit ...]")
    assert "truncated" in result
