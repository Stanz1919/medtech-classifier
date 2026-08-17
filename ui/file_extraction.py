"""Side-effect-free text extraction from uploaded files.

Turns a PDF, DOCX, plain-text, or image file into a plain string, so it
can be merged into the same free-text description box the user can type
into directly - the extraction/classification pipeline downstream never
knows or cares whether the text it received was typed, pasted, or pulled
from a file. No Streamlit import here on purpose (same reasoning as
``ui/examples.py``: keep anything that isn't a real Streamlit script
safely plain-importable and independently testable).

Image extraction is OCR only (pytesseract/Tesseract), by deliberate
choice - see the root README's "UI redesign" section for the full
reasoning. It reads text that is literally printed/rendered in the image
(dimension callouts, material labels, part numbers on a technical
drawing). It does NOT attempt to interpret shapes, form factor, or a
device's appearance in a photograph - doing that would mean routing the
image through a vision model's own judgement, which conflicts with this
project's whole premise ("not an LLM guessing tool"). A bare photo of a
device with no visible text will correctly, honestly yield nothing.
"""

from __future__ import annotations

import io
import os

MAX_EXTRACTED_CHARS = 4000

SUPPORTED_DOCUMENT_EXTENSIONS = ["pdf", "docx", "txt"]
SUPPORTED_IMAGE_EXTENSIONS = ["png", "jpg", "jpeg"]

# Tesseract's Windows installer does not add it to PATH. On Streamlit
# Community Cloud (Debian/apt via packages.txt) and most Linux/macOS
# installs it lands on PATH already, so this list is only ever consulted
# as a fallback after the PATH lookup fails.
_COMMON_WINDOWS_TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
]


class UnsupportedFileTypeError(ValueError):
    """Raised for a file extension none of the extractors handle."""


class TesseractNotAvailableError(RuntimeError):
    """Raised when OCR is requested but Tesseract isn't installed/found."""


class ImageDecodeError(ValueError):
    """Raised when an uploaded 'image' file can't actually be opened as one."""


def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) > MAX_EXTRACTED_CHARS:
        return text[:MAX_EXTRACTED_CHARS].rstrip() + "\n[... truncated - text was longer than the limit ...]"
    return text


def extract_text_from_txt(data: bytes) -> str:
    return _truncate(data.decode("utf-8", errors="replace"))


def extract_text_from_pdf(data: bytes) -> str:
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(data))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return _truncate("\n".join(pages_text))


def extract_text_from_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    # Technical files often put key specifications (material, sterility,
    # intended use) in tables rather than prose paragraphs.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return _truncate("\n".join(parts))


def _ensure_tesseract_configured() -> None:
    import pytesseract

    try:
        pytesseract.get_tesseract_version()
        return  # already discoverable on PATH
    except Exception:
        pass
    for path in _COMMON_WINDOWS_TESSERACT_PATHS:
        if os.path.isfile(path):
            pytesseract.pytesseract.tesseract_cmd = path
            return
    # Leave unconfigured - the OCR call below will raise a clear,
    # catchable error rather than silently doing nothing.


def is_tesseract_available() -> bool:
    try:
        _ensure_tesseract_configured()
        import pytesseract

        pytesseract.get_tesseract_version()
        return True
    except Exception:
        return False


def extract_text_from_image(data: bytes) -> str:
    import pytesseract
    from PIL import Image, UnidentifiedImageError

    try:
        image = Image.open(io.BytesIO(data))
        image.load()
    except UnidentifiedImageError as exc:
        raise ImageDecodeError("Could not read this file as an image - is it a valid PNG/JPEG?") from exc

    _ensure_tesseract_configured()
    try:
        text = pytesseract.image_to_string(image)
    except pytesseract.TesseractNotFoundError as exc:
        raise TesseractNotAvailableError(
            "OCR isn't available in this environment (Tesseract not found). On "
            "Streamlit Community Cloud this installs automatically via "
            "packages.txt; running locally, install Tesseract OCR and make sure "
            "it's on PATH - or just type/paste the description instead."
        ) from exc
    return _truncate(text)


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """Dispatch to the right extractor based on the uploaded file's
    extension. Raises UnsupportedFileTypeError / TesseractNotAvailableError
    / ImageDecodeError on failure - callers are expected to catch and
    display these, not let them propagate as a raw traceback."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return extract_text_from_pdf(data)
    if ext == "docx":
        return extract_text_from_docx(data)
    if ext == "txt":
        return extract_text_from_txt(data)
    if ext in SUPPORTED_IMAGE_EXTENSIONS:
        return extract_text_from_image(data)
    raise UnsupportedFileTypeError(f"Unsupported file type: .{ext or '(none)'}")
